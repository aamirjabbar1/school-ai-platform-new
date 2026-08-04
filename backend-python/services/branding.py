"""
Official Lahore School System (LSS) document template.

This module is the single formatting standard for *every* academic document the
platform hands to a user — question papers, model/past papers, monthly, midterm
and final exams, MCQ sheets, assignments, worksheets, lesson plans, teacher
notes and classroom activities — in both PDF and Microsoft Word. Exporters never
build their own page furniture: they hand a story (PDF) or a Document (Word) to
the helpers here and the branded page frame, header, footer and copyright notice
are applied automatically.

The template is configuration, not per-document state. The school's assets live
in `backend-python/assets/` and are picked up on import:

    lss_logo.png       school crest drawn in the page header (required)
    lss_template.pdf   optional. When present and its page geometry matches the
                       document, generated pages are stamped straight onto the
                       official artwork instead of the drawn header/frame — the
                       most faithful reproduction of the supplied template.
    lss_wordmark.ttf   optional blackletter face for the header wordmark;
                       falls back to Times-Bold.

Every path can be overridden with LSS_LOGO_PATH / LSS_TEMPLATE_PDF /
LSS_WORDMARK_FONT, and the wording with LSS_SCHOOL_NAME / LSS_COPYRIGHT_NOTICE.
"""
from __future__ import annotations

import io
import os
from functools import lru_cache
from pathlib import Path

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas as pdfcanvas
from reportlab.platypus import SimpleDocTemplate

# ─── Configured template assets ───────────────────────────────────────────────

_ASSETS_DIR = Path(__file__).resolve().parent.parent / "assets"

LOGO_PATH = os.getenv("LSS_LOGO_PATH") or str(_ASSETS_DIR / "lss_logo.png")
TEMPLATE_PDF_PATH = os.getenv("LSS_TEMPLATE_PDF") or str(_ASSETS_DIR / "lss_template.pdf")
WORDMARK_FONT_PATH = os.getenv("LSS_WORDMARK_FONT") or str(_ASSETS_DIR / "lss_wordmark.ttf")

# Kept separate from settings.SCHOOL_NAME, which names the chat assistant and is
# deliberately informal ("LSS Bot"); documents always carry the legal name.
SCHOOL_NAME = os.getenv("LSS_SCHOOL_NAME") or "Lahore School System"
COPYRIGHT_NOTICE = os.getenv("LSS_COPYRIGHT_NOTICE") or (
    "All rights reserved. Unconstitutional reproduction of this documents shall "
    "be liable for prosecution under the copy rights act."
)

# ─── Page geometry (shared by the PDF and Word renderers) ─────────────────────

BORDER_INSET = 1.0 * cm      # ruled box, measured in from every page edge
_LOGO_HEIGHT = 1.5 * cm
_LOGO_LEFT_PAD = 0.5 * cm
_LOGO_TOP_PAD = 0.4 * cm
_WORDMARK_SIZE = 19

_NOTICE_SIZE = 6.8
_NOTICE_BASELINE = BORDER_INSET + 0.42 * cm
_PAGENO_SIZE = 7
_PAGENO_BASELINE = BORDER_INSET + 0.95 * cm

LEFT_MARGIN = RIGHT_MARGIN = 1.8 * cm
TOP_MARGIN = 3.15 * cm       # clears the crest + wordmark
BOTTOM_MARGIN = 2.4 * cm     # clears the page number + copyright notice

PAGE_MARGINS = {
    "leftMargin": LEFT_MARGIN,
    "rightMargin": RIGHT_MARGIN,
    "topMargin": TOP_MARGIN,
    "bottomMargin": BOTTOM_MARGIN,
}


def content_width(pagesize=A4) -> float:
    """Usable width between the branded margins — use this to size tables."""
    return pagesize[0] - LEFT_MARGIN - RIGHT_MARGIN


# ─── Cached asset loading ─────────────────────────────────────────────────────

@lru_cache(maxsize=1)
def _logo() -> ImageReader | None:
    path = Path(LOGO_PATH)
    if not path.is_file():
        print(f"[branding] logo not found at {path} — header will omit the crest")
        return None
    try:
        return ImageReader(str(path))
    except Exception as exc:
        print(f"[branding] could not load logo {path}: {exc}")
        return None


@lru_cache(maxsize=1)
def _wordmark_font() -> str:
    """Register the school's blackletter face if one was configured."""
    path = Path(WORDMARK_FONT_PATH)
    if path.is_file():
        try:
            pdfmetrics.registerFont(TTFont("LSSWordmark", str(path)))
            return "LSSWordmark"
        except Exception as exc:
            print(f"[branding] could not register wordmark font {path}: {exc}")
    return "Times-Bold"


# ─── Document typeface ────────────────────────────────────────────────────────
#
# reportlab's built-in faces are Latin-1 only, so ℝ, ⁻¹, vector notation and
# other maths and science glyphs print as hollow boxes — unusable on a Physics
# or Mathematics paper. DejaVu Sans covers them and ships with the image
# (fonts-dejavu-core). If it is ever missing we fall back to Helvetica rather
# than fail: text stays readable, only exotic glyphs suffer.

_DEJAVU_DIR = Path(os.getenv("LSS_FONT_DIR") or "/usr/share/fonts/truetype/dejavu")
DOCUMENT_FONT_FAMILY = "LSSBody"

_FALLBACK_FONTS = {
    "regular": "Helvetica", "bold": "Helvetica-Bold",
    "italic": "Helvetica-Oblique", "bold_italic": "Helvetica-BoldOblique",
}


@lru_cache(maxsize=1)
def document_fonts() -> dict[str, str]:
    """Registered font names for the document body, by weight/style."""
    regular = _DEJAVU_DIR / "DejaVuSans.ttf"
    bold = _DEJAVU_DIR / "DejaVuSans-Bold.ttf"
    if not (regular.is_file() and bold.is_file()):
        print(f"[branding] DejaVu fonts not found in {_DEJAVU_DIR} — falling back to Helvetica; "
              "maths and science glyphs may not render")
        return dict(_FALLBACK_FONTS)

    try:
        pdfmetrics.registerFont(TTFont(DOCUMENT_FONT_FAMILY, str(regular)))
        pdfmetrics.registerFont(TTFont(f"{DOCUMENT_FONT_FAMILY}-Bold", str(bold)))
        # fonts-dejavu-core ships no oblique sans face; slanted text falls back
        # to the upright cut, which reads better than a missing font.
        pdfmetrics.registerFontFamily(
            DOCUMENT_FONT_FAMILY,
            normal=DOCUMENT_FONT_FAMILY,
            bold=f"{DOCUMENT_FONT_FAMILY}-Bold",
            italic=DOCUMENT_FONT_FAMILY,
            boldItalic=f"{DOCUMENT_FONT_FAMILY}-Bold",
        )
    except Exception as exc:
        print(f"[branding] could not register document fonts: {exc}")
        return dict(_FALLBACK_FONTS)

    return {
        "regular": DOCUMENT_FONT_FAMILY,
        "bold": f"{DOCUMENT_FONT_FAMILY}-Bold",
        "italic": DOCUMENT_FONT_FAMILY,
        "bold_italic": f"{DOCUMENT_FONT_FAMILY}-Bold",
    }


_BUILTIN_TO_ROLE = {
    "Helvetica": "regular",         "Times-Roman": "regular",
    "Helvetica-Bold": "bold",       "Times-Bold": "bold",
    "Helvetica-Oblique": "italic",  "Times-Italic": "italic",
    "Helvetica-BoldOblique": "bold_italic",
}


def use_document_fonts(styles: dict) -> dict:
    """Point a stylesheet at the Unicode document family, in place."""
    fonts = document_fonts()
    for style in styles.values():
        for attr in ("fontName", "bulletFontName"):
            role = _BUILTIN_TO_ROLE.get(getattr(style, attr, None))
            if role:
                setattr(style, attr, fonts[role])
    return styles


@lru_cache(maxsize=1)
def _template_pdf() -> bytes | None:
    """Raw bytes of the official template PDF, if the school supplied one."""
    path = Path(TEMPLATE_PDF_PATH)
    if not path.is_file():
        return None
    try:
        return path.read_bytes()
    except Exception as exc:
        print(f"[branding] could not read template {path}: {exc}")
        return None


@lru_cache(maxsize=1)
def _template_pagesize() -> tuple[float, float] | None:
    data = _template_pdf()
    if not data:
        return None
    try:
        from pypdf import PdfReader
        box = PdfReader(io.BytesIO(data)).pages[0].mediabox
        return float(box.width), float(box.height)
    except Exception as exc:
        print(f"[branding] could not read template page size: {exc}")
        return None


def _template_matches(pagesize) -> bool:
    """The artwork can only be stamped onto pages of the same geometry."""
    size = _template_pagesize()
    if not size:
        return False
    return abs(size[0] - pagesize[0]) < 2 and abs(size[1] - pagesize[1]) < 2


# ─── Drawn page furniture ─────────────────────────────────────────────────────

def _draw_frame(canvas, width: float, height: float) -> None:
    canvas.setStrokeColor(colors.black)
    canvas.setLineWidth(0.8)
    canvas.rect(
        BORDER_INSET, BORDER_INSET,
        width - 2 * BORDER_INSET, height - 2 * BORDER_INSET,
        stroke=1, fill=0,
    )


def _draw_header(canvas, width: float, height: float) -> None:
    logo_top = height - BORDER_INSET - _LOGO_TOP_PAD
    logo_bottom = logo_top - _LOGO_HEIGHT

    logo = _logo()
    if logo is not None:
        img_w, img_h = logo.getSize()
        draw_w = _LOGO_HEIGHT * (img_w / img_h) if img_h else _LOGO_HEIGHT
        canvas.drawImage(
            logo, BORDER_INSET + _LOGO_LEFT_PAD, logo_bottom,
            width=draw_w, height=_LOGO_HEIGHT,
            preserveAspectRatio=True, anchor="sw", mask="auto",
        )

    canvas.setFillColor(colors.black)
    canvas.setFont(_wordmark_font(), _WORDMARK_SIZE)
    canvas.drawCentredString(width / 2, (logo_top + logo_bottom) / 2 - 6, SCHOOL_NAME)


def _draw_notice(canvas, width: float) -> None:
    canvas.setFillColor(colors.HexColor("#111111"))
    canvas.setFont("Helvetica", _NOTICE_SIZE)
    canvas.drawCentredString(width / 2, _NOTICE_BASELINE, COPYRIGHT_NOTICE)


def _draw_page_number(canvas, width: float, page: int, total: int) -> None:
    canvas.setFillColor(colors.HexColor("#4b5563"))
    canvas.setFont("Helvetica", _PAGENO_SIZE)
    canvas.drawRightString(
        width - BORDER_INSET - _LOGO_LEFT_PAD, _PAGENO_BASELINE,
        f"Page {page} of {total}",
    )


def _make_decorator(*, draw_furniture: bool):
    """Per-page callback for SimpleDocTemplate.

    When the official template artwork is being stamped underneath, the frame,
    crest and notice already come from that artwork and must not be drawn twice.
    """
    def decorate(canvas, doc):
        if not draw_furniture:
            return
        canvas.saveState()
        width, height = canvas._pagesize
        _draw_frame(canvas, width, height)
        _draw_header(canvas, width, height)
        _draw_notice(canvas, width)
        canvas.restoreState()

    return decorate


class _NumberedCanvas(pdfcanvas.Canvas):
    """Two-pass canvas so every page can print "Page N of TOTAL"."""

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_pages: list[dict] = []

    def showPage(self):
        self._saved_pages.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        total = len(self._saved_pages)
        for index, state in enumerate(self._saved_pages, start=1):
            self.__dict__.update(state)
            self.saveState()
            _draw_page_number(self, self._pagesize[0], index, total)
            self.restoreState()
            super().showPage()
        super().save()


# ─── Public PDF entry point ───────────────────────────────────────────────────

def build_branded_pdf(story: list, *, pagesize=A4, title: str = "Document") -> bytes:
    """Render a platypus story inside the official LSS template."""
    stamp = _template_matches(pagesize)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=pagesize, title=title, **PAGE_MARGINS)
    decorate = _make_decorator(draw_furniture=not stamp)
    doc.build(story, onFirstPage=decorate, onLaterPages=decorate,
              canvasmaker=_NumberedCanvas)

    data = buf.getvalue()
    return _stamp_template(data) if stamp else data


def _stamp_template(pdf_bytes: bytes) -> bytes:
    """Lay each generated page over the official template artwork."""
    template = _template_pdf()
    if not template:
        return pdf_bytes
    try:
        from pypdf import PdfReader, PdfWriter

        writer = PdfWriter()
        for page in PdfReader(io.BytesIO(pdf_bytes)).pages:
            # Re-read per page: merge_page mutates the base, so each page needs
            # its own pristine copy of the artwork.
            base = PdfReader(io.BytesIO(template)).pages[0]
            base.merge_page(page)
            writer.add_page(base)
        out = io.BytesIO()
        writer.write(out)
        return out.getvalue()
    except Exception as exc:
        # Never lose the document over a branding problem — the drawn header is
        # already on the page in that case.
        print(f"[branding] template stamping failed, returning plain render: {exc}")
        return pdf_bytes


# ─── Public Word (.docx) entry point ──────────────────────────────────────────

_SECTPR_SUCCESSORS = (
    "lnNumType", "pgNumType", "cols", "formProt", "vAlign", "noEndnote",
    "titlePg", "textDirection", "bidi", "rtlGutter", "docGrid",
    "printerSettings", "sectPrChange",
)


def apply_docx_branding(document) -> None:
    """Apply the LSS margins, ruled border, header and footer to every section.

    Call this *after* setting page size and orientation — the header table is
    sized from the section's own geometry.
    """
    from docx.shared import Cm

    for section in document.sections:
        section.left_margin = Cm(LEFT_MARGIN / cm)
        section.right_margin = Cm(RIGHT_MARGIN / cm)
        section.top_margin = Cm(TOP_MARGIN / cm)
        section.bottom_margin = Cm(BOTTOM_MARGIN / cm)
        # Keep header/footer content clear of the ruled border.
        section.header_distance = Cm((BORDER_INSET + 0.40 * cm) / cm)
        section.footer_distance = Cm((BORDER_INSET + 0.35 * cm) / cm)
        section.different_first_page_header_footer = False

        _docx_page_border(section)
        _docx_header(section)
        _docx_footer(section)


def _oxml(tag: str):
    from docx.oxml import OxmlElement
    return OxmlElement(tag)


def _reset_container(container) -> None:
    """Empty a header/footer down to a single blank paragraph.

    Clearing runs in place (rather than assigning `.text = ""`) avoids leaving
    a stray empty run behind, and keeps the call idempotent.
    """
    for para in list(container.paragraphs)[1:]:
        para._p.getparent().remove(para._p)
    for table in list(container.tables):
        table._tbl.getparent().remove(table._tbl)
    if not container.paragraphs:
        container.add_paragraph()
    for run in list(container.paragraphs[0].runs):
        run._r.getparent().remove(run._r)


def _docx_page_border(section) -> None:
    """Draw the ruled box, measured from the page edge like the PDF frame."""
    from docx.oxml.ns import qn

    sect_pr = section._sectPr
    for existing in sect_pr.findall(qn("w:pgBorders")):
        sect_pr.remove(existing)

    # w:space is whole points from the page edge and Word caps it at 31.
    inset_pt = min(31, max(0, round(BORDER_INSET)))

    borders = _oxml("w:pgBorders")
    borders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        el = _oxml(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")            # eighths of a point → 0.75pt
        el.set(qn("w:space"), str(inset_pt))
        el.set(qn("w:color"), "000000")
        borders.append(el)

    # w:pgBorders has a fixed position in the sectPr schema; append would land it
    # after w:cols/w:docGrid and Word flags the file as needing repair.
    anchor = None
    for tag in _SECTPR_SUCCESSORS:
        found = sect_pr.find(qn(f"w:{tag}"))
        if found is not None:
            anchor = found
            break
    if anchor is not None:
        anchor.addprevious(borders)
    else:
        sect_pr.append(borders)


def _docx_header(section) -> None:
    """Crest on the left, wordmark centred — mirrors the PDF header."""
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    header = section.header
    header.is_linked_to_previous = False
    _reset_container(header)

    usable = section.page_width - section.left_margin - section.right_margin
    logo_col = Cm(3.0)
    name_col = usable - 2 * logo_col

    # Three columns, with an empty one mirroring the crest, so the wordmark
    # centres on the page rather than on the space left of the crest — matching
    # where build_branded_pdf draws it.
    table = header.add_table(rows=1, cols=3, width=usable)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # Column widths drive w:gridCol, cell widths drive w:tcW — Word needs both
    # to agree or a fixed-layout table renders at the wrong proportions.
    widths = (logo_col, name_col, logo_col)
    for column, width in zip(table.columns, widths):
        column.width = width

    logo_cell, name_cell, _spacer = table.rows[0].cells
    for cell, width in zip(table.rows[0].cells, widths):
        cell.width = width

    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_para.paragraph_format.space_after = Pt(0)
    if Path(LOGO_PATH).is_file():
        try:
            logo_para.add_run().add_picture(LOGO_PATH, height=Cm(_LOGO_HEIGHT / cm))
        except Exception as exc:
            print(f"[branding] could not embed logo in Word header: {exc}")

    name_para = name_cell.paragraphs[0]
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(0)
    name_run = name_para.add_run(SCHOOL_NAME)
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # A table must be followed by a paragraph, so reuse the header's own empty
    # one — moved after the table and collapsed so it adds no visible height
    # (an ordinary empty line here would push the body past the top margin).
    trailing = header.paragraphs[0]
    trailing.paragraph_format.space_before = Pt(0)
    trailing.paragraph_format.space_after = Pt(0)
    trailing.paragraph_format.line_spacing = Pt(1)
    trailing._p.addprevious(table._tbl)


def _docx_footer(section) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    footer = section.footer
    footer.is_linked_to_previous = False
    _reset_container(footer)

    page_para = footer.paragraphs[0]
    page_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    page_para.paragraph_format.space_before = Pt(0)
    page_para.paragraph_format.space_after = Pt(0)
    _footer_run(page_para, "Page ", size=_PAGENO_SIZE, color=(0x4B, 0x55, 0x63))
    _docx_field(page_para, "PAGE", size=_PAGENO_SIZE, color=(0x4B, 0x55, 0x63))
    _footer_run(page_para, " of ", size=_PAGENO_SIZE, color=(0x4B, 0x55, 0x63))
    _docx_field(page_para, "NUMPAGES", size=_PAGENO_SIZE, color=(0x4B, 0x55, 0x63))

    notice_para = footer.add_paragraph()
    notice_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice_para.paragraph_format.space_before = Pt(0)
    notice_para.paragraph_format.space_after = Pt(0)
    _footer_run(notice_para, COPYRIGHT_NOTICE, size=_NOTICE_SIZE, color=(0x11, 0x11, 0x11))


def _footer_run(paragraph, text: str, *, size: float, color: tuple[int, int, int]):
    from docx.shared import Pt, RGBColor

    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    return run


def _docx_field(paragraph, instruction: str, *, size: float, color: tuple[int, int, int]) -> None:
    """Insert a Word field (PAGE / NUMPAGES) with a cached placeholder result.

    Each part gets its own run carrying the footer formatting: when Word or
    LibreOffice recomputes the field it reuses the result run's properties, so
    a single combined run would repaint the number at the default body size.
    """
    from docx.oxml.ns import qn

    def part(node):
        _footer_run(paragraph, "", size=size, color=color)._r.append(node)

    begin = _oxml("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    part(begin)

    instr = _oxml("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    part(instr)

    separate = _oxml("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    part(separate)

    _footer_run(paragraph, "1", size=size, color=color)  # cached result

    end = _oxml("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    part(end)
