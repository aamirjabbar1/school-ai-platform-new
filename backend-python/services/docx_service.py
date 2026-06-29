"""
Render a lesson plan into a Microsoft Word (.docx) document using python-docx.

    build_lesson_plan_docx(plan) -> bytes

Produces a landscape document with a heading, an optional overview, a schedule
table (one row per lesson), and a summary table — mirroring the PDF export so
teachers can edit the plan in Word.
"""
from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Mirrors _PLAN_CELL_GROUPS in pdf_service so both exports show the same fields.
_PLAN_CELL_GROUPS = [
    ("Topic", [("chapter", "Chapter"), ("topic", "Topic"), ("subtopic", "Subtopic")]),
    ("Objectives & Outcomes", [("objectives", "Objectives"), ("outcomes", "Outcomes"), ("prior_knowledge", "Prior knowledge")]),
    ("Methodology & Activities", [("methodology", "Method"), ("teacher_activities", "Teacher"), ("student_activities", "Students"), ("group_activity", "Activity"), ("hots", "HOTS")]),
    ("Resources & Examples", [("resources", "Resources"), ("real_life_examples", "Real-life"), ("differentiation", "Differentiation")]),
    ("Homework & Assessment", [("classwork", "Classwork"), ("homework", "Homework"), ("assessment", "Assessment"), ("remarks", "Remarks")]),
]

_SUMMARY_LABELS = [
    ("academic_session", "Academic Session"), ("subject", "Subject"),
    ("grade", "Grade"), ("board", "Board"),
    ("total_chapters", "Total Chapters"), ("total_weeks", "Total Teaching Weeks"),
    ("total_teaching_days", "Total Teaching Days"), ("total_lessons", "Total Lessons"),
    ("total_practical_lessons", "Total Practical Lessons"),
    ("total_assessments", "Total Assessments"), ("total_homework", "Total Homework"),
    ("revision_schedule", "Revision Schedule"),
    ("exam_prep_schedule", "Exam Preparation Schedule"),
    ("expected_completion_date", "Expected Completion Date"),
]

_HEADER_BG = "1E3A8A"


def _shade_cell(cell, hex_color: str) -> None:
    """Apply a solid background fill to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_text(cell, runs: list[tuple[str, str]], *, size: float = 7.5,
                   white: bool = False, bold_all: bool = False) -> None:
    """Fill a cell with labelled fragments. `runs` is a list of (label, value);
    an empty label renders the value plainly."""
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    first = True
    for label, value in runs:
        if not value and not label:
            continue
        if not first:
            para.add_run().add_break()
        first = False
        if label:
            r = para.add_run(f"{label}: ")
            r.bold = True
            r.font.size = Pt(size)
            if white:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r2 = para.add_run(str(value))
        r2.font.size = Pt(size)
        r2.bold = bold_all
        if white:
            r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if first:  # nothing written — keep an empty run so the cell isn't collapsed
        para.add_run("")


def build_lesson_plan_docx(plan: dict) -> bytes:
    plan_data = plan.get("plan_data") or {}
    lessons = plan_data.get("lessons") or []
    summary = plan_data.get("summary") or {}

    doc = Document()

    # Landscape A4
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = Cm(29.7), Cm(21.0)
    section.left_margin = section.right_margin = Cm(1.2)
    section.top_margin = section.bottom_margin = Cm(1.2)

    # Title + meta
    title = doc.add_heading(plan.get("title", "Lesson Plan"), level=0)
    title.alignment = WD_TABLE_ALIGNMENT.CENTER

    meta_bits = [plan.get("subject"), plan.get("class_name")]
    if plan.get("section"):
        meta_bits.append(f"Section {plan['section']}")
    if plan.get("board"):
        meta_bits.append(plan["board"])
    ptype = (plan.get("plan_type") or "").replace("_", " ").title()
    if ptype:
        meta_bits.append(f"{ptype} Plan")
    if plan.get("academic_session"):
        meta_bits.append(f"Session {plan['academic_session']}")
    meta_para = doc.add_paragraph(" • ".join(str(b) for b in meta_bits if b))
    meta_para.alignment = WD_TABLE_ALIGNMENT.CENTER
    for run in meta_para.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    overview = plan_data.get("overview")
    if overview:
        doc.add_paragraph(overview)

    # Schedule table
    headers = ["Wk / Date"] + [g[0] for g in _PLAN_CELL_GROUPS]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _shade_cell(hdr[i], _HEADER_BG)
        _set_cell_text(hdr[i], [("", h)], size=8.5, white=True, bold_all=True)

    for ls in lessons:
        cells = table.add_row().cells
        wk_runs: list[tuple[str, str]] = []
        if ls.get("week") is not None:
            wk_runs.append(("", f"Week {ls.get('week')}"))
        for key in ("dates", "day", "period"):
            if ls.get(key):
                wk_runs.append(("", ls.get(key)))
        _set_cell_text(cells[0], wk_runs or [("", "")], size=8.5, bold_all=True)
        for ci, (_, fields) in enumerate(_PLAN_CELL_GROUPS, start=1):
            runs = [(label, ls.get(key)) for key, label in fields if ls.get(key)]
            _set_cell_text(cells[ci], runs or [("", "")], size=7.5)

    # Summary
    rows = [(label, summary.get(key)) for key, label in _SUMMARY_LABELS
            if summary.get(key) not in (None, "")]
    if rows:
        doc.add_heading("Plan Summary", level=1)
        sum_tbl = doc.add_table(rows=0, cols=2)
        sum_tbl.style = "Table Grid"
        for label, value in rows:
            c = sum_tbl.add_row().cells
            _shade_cell(c[0], "F3F4F6")
            _set_cell_text(c[0], [("", label)], size=9.5, bold_all=True)
            _set_cell_text(c[1], [("", value)], size=9.5)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
